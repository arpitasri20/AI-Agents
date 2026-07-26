"""
RESUME / JD MATCHER — Streamlit UI
======================================
Paste any job description and resume text, get a structured fit
analysis. No external data API dependency — just Gemini + local logic.
"""

import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field
from typing import List
import io
from pypdf import PdfReader
from docx import Document


def extract_text_from_file(uploaded_file) -> str:
    """Extract plain text from an uploaded PDF, DOCX, or TXT file."""
    if uploaded_file is None:
        return ""

    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        # Some documents (especially content pasted from web pages) store
        # text inside tables rather than plain paragraphs — paragraphs-only
        # extraction misses this entirely, so we check tables too.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    elif filename.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")

    else:
        return ""


def input_section(label: str, key_prefix: str) -> str:
    """Renders a toggle between pasting text and uploading a file for
    one input (JD or resume), returning whichever text is available."""
    input_mode = st.radio(
        f"{label} input method",
        ["Paste text", "Upload file"],
        key=f"{key_prefix}_mode",
        horizontal=True,
        label_visibility="collapsed",
    )

    if input_mode == "Paste text":
        return st.text_area(
            label,
            height=280,
            placeholder=f"Paste the {label.lower()} text here...",
            key=f"{key_prefix}_text",
        )
    else:
        uploaded = st.file_uploader(
            f"Upload {label} (PDF, DOCX, or TXT)",
            type=["pdf", "docx", "txt"],
            key=f"{key_prefix}_file",
        )
        extracted = extract_text_from_file(uploaded)
        if extracted:
            st.success(f"Extracted {len(extracted)} characters from {uploaded.name}")
            with st.expander("Preview extracted text"):
                st.text(extracted[:1000] + ("..." if len(extracted) > 1000 else ""))
        return extracted


class MatchAnalysis(BaseModel):
    match_score: int = Field(description="Overall fit score from 0 to 100")
    matched_skills: List[str] = Field(
        description="Skills/requirements from the JD that the resume clearly demonstrates"
    )
    missing_skills: List[str] = Field(
        description="Skills/requirements from the JD NOT evident in the resume"
    )
    suggested_resume_tweaks: List[str] = Field(
        description="3-5 specific, actionable suggestions to improve fit"
    )
    overall_verdict: str = Field(
        description="One or two sentence honest verdict on fit"
    )


match_prompt = ChatPromptTemplate.from_template(
    """You are an experienced technical recruiter evaluating fit between a candidate's resume and a job description.

Be specific and honest — do not inflate the score to be encouraging. Base matched_skills and missing_skills
on what is ACTUALLY written in the resume, not what the candidate might plausibly know.

JOB DESCRIPTION:
{job_description}

RESUME:
{resume_text}

Respond ONLY as JSON matching this schema:
{{
  "match_score": int,
  "matched_skills": [str, ...],
  "missing_skills": [str, ...],
  "suggested_resume_tweaks": [str, ...],
  "overall_verdict": str
}}"""
)


@st.cache_resource
def build_chain():
    llm = ChatGoogleGenerativeAI(
        model="gemini-flash-latest",
        google_api_key=st.secrets.get("GEMINI_API_KEY", None),
    )
    return match_prompt | llm | JsonOutputParser(pydantic_object=MatchAnalysis)


# ---------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------
st.set_page_config(page_title="Resume ↔ JD Matcher", page_icon="📄")
st.title("📄 Resume ↔ Job Description Matcher")
st.caption(
    "Paste a job description and your resume text to see your fit — score, matched skills, gaps, and specific tweaks."
)

col1, col2 = st.columns(2)
with col1:
    st.markdown("### Job Description")
    jd_text = input_section("Job Description", "jd")
with col2:
    st.markdown("### Your Resume")
    resume_text = input_section("Resume", "resume")

analyze_clicked = st.button("Analyze Fit", type="primary")

if analyze_clicked:
    if not jd_text or not resume_text:
        st.warning("Please paste both the job description and your resume.")
    else:
        chain = build_chain()
        with st.spinner("Analyzing fit..."):
            result = chain.invoke(
                {"job_description": jd_text, "resume_text": resume_text}
            )

        st.divider()

        # Score with a visual progress bar
        score = result["match_score"]
        st.subheader(f"Match Score: {score}/100")
        st.progress(score / 100)

        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("**✅ Matched Skills**")
            for skill in result["matched_skills"]:
                st.write(f"- {skill}")
        with col_b:
            st.markdown("**❌ Missing Skills**")
            for skill in result["missing_skills"]:
                st.write(f"- {skill}")

        st.markdown("**💡 Suggested Resume Tweaks**")
        for tweak in result["suggested_resume_tweaks"]:
            st.write(f"- {tweak}")

        st.markdown("**📝 Verdict**")
        st.info(result["overall_verdict"])
